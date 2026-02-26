from fastapi import FastAPI, HTTPException
from pydantic import BaseModel, Field
from typing import List, Literal, Optional
from docx import Document
from fastapi.responses import FileResponse
import uuid
import os
import uvicorn

app = FastAPI()

TEMPLATE_DIR = "templates"

# ✅ Разрешаем только те стили, которые есть в твоём .docx шаблоне
ALLOWED_STYLES = {
    "GBM Normal1",
    "GBM Dash Bullet",
    "GBM Note",
    "GBM Raw",
    "Heading 1",
    "Heading 2",
    "Heading 3",
}

BlockType = Literal["heading", "paragraph", "bullet_list", "note", "raw"]
HeadingLevel = Literal[1, 2, 3]


class StyledBlock(BaseModel):
    type: BlockType
    text: str = Field(..., description="Raw text, no summarization")
    word_style: str = Field(..., description="Must exist in template")
    level: Optional[HeadingLevel] = Field(
        default=None, description="Required only when type=heading"
    )


class BuilderPayload(BaseModel):
    title: Optional[str] = None
    document_type: Optional[str] = None
    styled_blocks: List[StyledBlock] = Field(default_factory=list)


class RequestBody(BaseModel):
    template: str
    payload: BuilderPayload


@app.get("/health")
def health():
    return {"ok": True}


@app.post("/")
def generate_docx(req: RequestBody):
    # --- Validate template file exists ---
    template_path = os.path.join(TEMPLATE_DIR, req.template)
    if not os.path.exists(template_path):
        raise HTTPException(status_code=404, detail=f"Template not found: {req.template}")

    # --- Validate blocks present ---
    blocks = req.payload.styled_blocks
    if not blocks or len(blocks) == 0:
        raise HTTPException(status_code=400, detail="payload.styled_blocks is required and cannot be empty")

    # --- Validate every block against strict rules ---
    for i, b in enumerate(blocks):
        # style whitelist
        if b.word_style not in ALLOWED_STYLES:
            raise HTTPException(
                status_code=400,
                detail=f"Block #{i} has unsupported word_style='{b.word_style}'. Allowed: {sorted(ALLOWED_STYLES)}",
            )

        # heading level rules
        if b.type == "heading":
            if b.level not in (1, 2, 3):
                raise HTTPException(
                    status_code=400,
                    detail=f"Block #{i} type='heading' requires level=1|2|3",
                )
        else:
            if b.level is not None:
                raise HTTPException(
                    status_code=400,
                    detail=f"Block #{i} type='{b.type}' must NOT include 'level'",
                )

        # text must be string (non-null); empty allowed but usually useless
        if b.text is None:
            raise HTTPException(status_code=400, detail=f"Block #{i} text must not be null")

    # --- Render DOCX using template ---
    doc = Document(template_path)

    # Optional: Add title as Heading 1 only if provided
    if req.payload.title and req.payload.title.strip():
        # Use Heading 1 to keep template consistency
        doc.add_paragraph(req.payload.title.strip(), style="Heading 1")

    # Optional: Add document_type as Note (or Normal) if provided
    if req.payload.document_type and req.payload.document_type.strip():
        # If your template has a standard place for document type, this keeps it consistent
        # You can change style to "GBM Note" or "GBM Normal1"
        doc.add_paragraph(req.payload.document_type.strip(), style="GBM Note")

    # Render blocks strictly
    for b in blocks:
        if b.type == "heading":
            # We enforce Heading 1/2/3 via style mapping
            # and level drives which heading style to use.
            heading_style = f"Heading {b.level}"
            doc.add_paragraph(b.text, style=heading_style)

        elif b.type == "bullet_list":
            # Each line is a bullet item
            for line in b.text.split("\n"):
                line = line.strip()
                if line:
                    doc.add_paragraph(line, style=b.word_style)

        else:
            # paragraph / note / raw
            doc.add_paragraph(b.text, style=b.word_style)

    # Save to /tmp (Cloud Run writable)
    out_path = f"/tmp/output_{uuid.uuid4()}.docx"
    doc.save(out_path)

    return FileResponse(
        out_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename="formatted_document.docx",
    )


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 8080))
    uvicorn.run(app, host="0.0.0.0", port=port)
