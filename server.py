from flask import Flask, request, jsonify
from docx import Document
import tempfile
import base64

app = Flask(__name__)

@app.post("/")
def build_docx():
    data = request.json
    doc = Document()

    # Title
    if "title" in data:
        doc.add_heading(data["title"], level=1)

    # Content blocks
    for block in data.get("styled_blocks", []):
        if block["type"] == "heading":
            doc.add_heading(block["text"], level=block.get("level", 2))
        elif block["type"] == "paragraph":
            doc.add_paragraph(block["text"])
        elif block["type"] == "bullet_list":
            for line in block["text"].split("\n"):
                doc.add_paragraph(line.replace("-", "").strip(), style="List Bullet")

    # Save to temp file
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(tmp.name)

    with open(tmp.name, "rb") as f:
        encoded = base64.b64encode(f.read()).decode()

    return jsonify({"docx_base64": encoded})
